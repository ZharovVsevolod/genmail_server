from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from datetime import datetime
from random import randint

from gm_services.neural.llm import LLModelShell
from gm_services.database.tablestore import PGHandler
from gm_services.database.tablestore.table_schemas.user import User
from gm_services.schemas.understanding import DocumentView
from gm_services.common import (
    generate_hex, 
    load_prompt, 
    simple_langchain_messages
)
from gm_services.config import Settings

from docx.document import Document as DocumentObject
from typing import Literal

import logging
logger = logging.getLogger(__name__)


DOCUMENT_TYPE = Literal["docx", "txt"]

INNER_DOCUMENT_BLANK = Settings.docs.full_mail_path + "blanks/outer.docx"
OUTER_DOCUMENT_BLANK = Settings.docs.full_mail_path + "blanks/outer.docx"


class DocumentFormalizer:
    def __init__(
        self, 
        model: LLModelShell,
        tablestore: PGHandler,
        path_to_save_folder: str | None = None
    ):
        # None check
        if path_to_save_folder is None:
            path_to_save_folder = Settings.docs.full_mail_path
        
        self.path_to_save_folder = path_to_save_folder

        # Links to outer modules
        self.model = model
        self.tablestore = tablestore
    

    def _load_document(self, document_info: DocumentView) -> DocumentObject:
        match document_info.doc_type:
            case "inner":
                path_to_load = INNER_DOCUMENT_BLANK

            case "outer":
                path_to_load = OUTER_DOCUMENT_BLANK
                
        document = Document(path_to_load)
        return document
    

    def _extract_metadata(self, document_info: DocumentView) -> DocumentView:
        logger.info("Get some information of document's author")
        # Get author information
        ## Nominative form
        logger.info("Generate Nominative form")
        system = load_prompt("nominative_form.md", default_path = True)
        messages = simple_langchain_messages(system, document_info.author)
        person_nominative = self.model.llm_answer(messages, llm_answer_parser = "json")
        
        ## Dative form
        logger.info("Generate Dative form")
        system = load_prompt("dative_form.md", default_path = True)
        human = f"{person_nominative["name"]}"
        human += f"{person_nominative["surname"]}"
        human += f"{person_nominative["patronymic"]}, "
        # ------------------------
        human += f"<ПОЗИЦИЯ>, "
        human += f"<ОРГАНИЗАЦИЯ>"
        # ------------------------
        messages = simple_langchain_messages(system, human)
        dative_form = self.model.llm_answer(messages, llm_answer_parser = "json")

        # Default values
        for key in dative_form.keys():
            if dative_form[key] is None:
                dative_form[key] = ""

        # Save all of it in document metadata
        full_dative_name = dative_form.get("name", "")
        full_dative_name += " " + dative_form.get("surname", "")
        full_dative_name += " " + dative_form.get("patronymic", "")
        document_info.metadata["author_position_dative"] = dative_form["position"]
        document_info.metadata["author_organization_genitive"] = dative_form["organization"]

        # Current date and number
        document_info.metadata["current_date"] = datetime.today().strftime("%d.%m.%Y")
        number = str(randint(1, 10000))
        number = "0" * (5 - len(number)) + number
        document_info.metadata["current_number"] = number

        return document_info


    def _generate_name(
            base_name: str | None = None, 
            doc_type: DOCUMENT_TYPE | None = None
        ) -> str:
        if base_name is not None:
            name = base_name.split(".")[0]
            doc_extension = base_name.split(".")[0]
        else:
            name = ""
            doc_extension = f".{doc_type}"
        
        name += generate_hex(8)
        name += doc_extension
        return name


    def _docx_head_pattern(
        self, 
        document: DocumentObject, 
        document_info: DocumentView
    ) -> DocumentObject:
        document.add_paragraph()
        document.add_paragraph()

        paragraph = document.add_paragraph(f"{document_info.metadata["author_position_dative"]}")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        paragraph = document.add_paragraph(f"{document_info.metadata["author_organization_genitive"]}")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph = document.add_paragraph(f"{document_info.metadata["author_dative"]}")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        document.add_paragraph()

        current_date_number = "      "
        current_date_number += str(document_info.metadata["current_date"])
        current_date_number += "                 "
        current_date_number += str(document_info.metadata["current_number"])
        p = document.add_paragraph(current_date_number)
        paragraph_format = p.paragraph_format
        paragraph_format.space_after = Pt(4.0)

        previous_number_date = "          "
        previous_number_date += document_info.number
        previous_number_date += "                      "
        previous_number_date += document_info.date
        p = document.add_paragraph(previous_number_date)

        document.add_paragraph()
        document.add_paragraph()
        document.add_paragraph()

        paragraph = document.add_paragraph(f"Уважаемый {document_info.author}!")
        document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return document


    def _docx_tail_pattern(self, document: DocumentObject, user: User) -> DocumentObject:
        document.add_paragraph()
        document.add_paragraph()
        document.add_paragraph("С уважением,")

        section = document.sections[0]
        page_width = section.page_width - section.left_margin - section.right_margin

        p = document.add_paragraph()

        # Configure paragraph tab stop at the right margin
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(page_width, alignment = 1)  # 1 = RIGHT

        # Add text: left text + tab + right text
        user_position = self.tablestore.find_user_position(user)
        p.add_run(user_position)
        p.add_run(f"\t{user.surname} {user.name} {user.patronymic}")

        return document


    def _docx_pattern(
        self, 
        user: User,
        document_info: DocumentView,
        body_text: str
    ) -> DocumentObject:
        # Load document depends on it's type
        document = self._load_document(document_info)
        # Generate essential metadate
        document_info = self._extract_metadata(document_info)

        # Head
        document = self._docx_head_pattern(document, document_info)
        # Actual text
        document.add_paragraph(body_text)
        # Tail
        document = self._docx_tail_pattern(document, user)

        return document

    
    def formalize(
        self,
        user: User,
        document_info: DocumentView,
        text: str,
        doc_type: DOCUMENT_TYPE = "docx",
        random_name: bool = False
    ) -> str:
        """
        Returns
        -------
        document_name: str
            Name of generated document
        """
        match doc_type:
            case "txt":
                path = self.path_to_save_folder + Settings.docs.final_name
                with open(path, "w+") as file:
                    file.write(text)
                return Settings.docs.final_name
            

            case "docx":
                document = self._docx_pattern(user, document_info, text)

                if random_name:
                    docname = self._generate_name(
                        base_name = Settings.docs.final_name,
                        doc_type = doc_type
                    )
                else:
                    docname = Settings.docs.final_name

                path_to_save = self.path_to_save_folder + docname
                
                document.save(path_to_save)
                return docname
